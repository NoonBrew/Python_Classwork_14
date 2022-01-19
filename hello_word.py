import docx

document = docx.Document()

document.add_paragraph('Hello Word!', 'Title')

document.add_paragraph('By Nate', 'Heading 1')

document.add_paragraph('This is a word document created with Python and python-docx')

document.add_paragraph('Automate the boring stuff.', 'Quote')

document.add_paragraph('This is the start of a list', 'List Bullet')

document.add_paragraph('List of favorite Colors', 'Heading 2')

favorite_colors = ['Blue', 'Purple', 'Orange']

for color in favorite_colors:
    document.add_paragraph(color, 'List Bullet 2')

document.save('Hello_word.docx')
