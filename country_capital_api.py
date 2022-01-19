from openpyxl import Workbook
import requests
import docx

# Initializes the word document object for the variable.
country_document = docx.Document()
# Initializes a exel spreadsheet object for the variable.
country_workbook = Workbook()
# Assigns the active worksheet of our workbook to the a variable.
country_worksheet = country_workbook.active
# Stores the URL of the API page we want to pull requests from.
url = 'https://country-list-1150.herokuapp.com/api/country'
# Requests the API page from the URL and stores it as a dictionary in a list.
country_dict = requests.get(url).json()
# Assigns a title to our worksheet that we will be writing data too.
country_worksheet.title = 'Country and Capitals'
# Writes 'Country' to row 1 column 1 of our worksheet.
country_worksheet.cell(1, 1, 'Country')
# Writes 'Capital City' to row 1 column 2 of our worksheet.
country_worksheet.cell(1, 2, 'Capital City')
# Unlike Python, Excel rows start counting at 1, we assign 2 since we have already writen information
# to row 1.
worksheet_index = 2
#  Loops through the list stored in country_dict and pulls the information from the dictionary for every entry.
for country in country_dict:
    # writes the name of each country to column 1 starting at row 2
    country_worksheet.cell(worksheet_index, 1, country['name'])
    # writes the capital city of each country to column 2 starting at row 2
    country_worksheet.cell(worksheet_index, 2, country['capitalCity'])
    # increasing the index so the loop can write to a new row.
    worksheet_index += 1

# creates a new paragraph in our word document with the text 'Countries of the World' and the heading style 1
country_document.add_paragraph('Countries of the World', 'Heading 1')
# Loops through our list for each entry and writes the name of the country and its capital city
for country in country_dict:
    # adds a new paragraph with the countries name in the heading 3 style
    country_document.add_paragraph(country['name'], 'Heading 3')
    # adds a new paragraph explaining what the capital city of each country is.
    country_document.add_paragraph(f'The capital city of {country["name"]} is {country["capitalCity"]}')
# saves our word document as a word file.
country_document.save('country_capitals.docx')
# saves our exel workbook as a excel file.
country_workbook.save('country_capitals.xlsx')