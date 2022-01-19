import requests, random, docx
from docx.enum.text import WD_BREAK

# Initializes a new word document. It is outside of our main function to act as a global variable.
state_park_brochure = docx.Document()


def word_doc_create(state_park_dict):
    """ Function that builds the word document. It takes a state park specific dictionary
    and then pulls information from that dictionary to copy to the word document."""

    # Initiates a list to hold the state park picture links.
    park_images = []
    # Access the park_images values of the state park dictionary and appends the picture links to our list
    for image in state_park_dict['park_images']:
        park_images.append(image)
    # picture saving code thanks to
    # https://www.kite.com/python/answers/how-to-download-an-image-using-requests-in-python
    # uses requests to download the first picture in our parks picture files to use as a header picture
    state_park_picture = requests.get(park_images[0])
    file = open('park_header_pic.png', 'wb')
    # saves the picture
    file.write(state_park_picture.content)

    # Pulls the name value from the state park dictionary and creates a new paragraph with the specified
    # heading.
    state_park_brochure.add_paragraph(state_park_dict['name'], 'Heading 1')
    # Inserts the header picture into the word document with a specified width
    state_park_brochure.add_picture('park_header_pic.png', width=docx.shared.Inches(5.5))
    # Inserts a heading for park highlights
    state_park_brochure.add_paragraph('Highlights', 'Heading 2')
    # access the state park dictionary highlight values and then creates a bullet for each highlight
    for highlight in state_park_dict['highlights']:
        state_park_brochure.add_paragraph(highlight, 'List Bullet')

    # Access the dictionary in stored in the park_information dictionary.
    for park_detail, text in state_park_dict['park_information'].items():
        # creates a new paragraph for each key stored under the park_information.
        state_park_brochure.add_paragraph(park_detail, 'Heading 2')
        # writes the value stored for each key contained in the park_information sub dictionary.
        state_park_brochure.add_paragraph(text)

    # uses a for loop to save than insert the remaining pictures in our picture list to the word document
    # we start at one so we do not repost the header picture and we leave the end blank so we don't have to know
    # the exact amount of picture links in our list.
    for image_link in park_images[1:]:
        # pulls the picture from the url link
        remaining_park_picture = requests.get(image_link)
        # opens the picture under the designated name
        picture_file = open('state_park_picture.png', 'wb')
        # saves our image to that file name.
        picture_file.write(remaining_park_picture.content)
        # inserts the image into the word document with a specified width.
        state_park_brochure.add_picture('state_park_picture.png', width=docx.shared.Inches(5))

    # Creates our contact information section of the word document.
    state_park_brochure.add_paragraph('Contact Information', 'Heading 2')
    state_park_brochure.add_paragraph('Address', 'Heading 3')
    # access the value stored under the address key of the dictionary and writes the information to the word doc.
    state_park_brochure.add_paragraph(state_park_dict['address'])

    state_park_brochure.add_paragraph('Website', 'Heading 2')
    # access the value stored under the url key of the dictionary and writes to the word doc.
    state_park_brochure.add_paragraph(state_park_dict['url'])


def mn_state_park_lookup(park_id_number):
    """Function that takes a state park ID number from a list and appends it to a URL to access
     a state park specific webpage with a state park specific dictionary."""
    state_park_id_url = "https://mn-state-parks.herokuapp.com/api/"
    # appends the id number to the base URL to pull the page request and store it in a dictionary.
    state_park_detail_dict = requests.get(state_park_id_url + park_id_number).json()
    # runs the dictonary through are word_doc_create function.
    word_doc_create(state_park_detail_dict)


def main():
    # variable that stores the URL of the API we that lists multiple state parks in Minnesota.
    state_park_list_url = "https://mn-state-parks.herokuapp.com/api/list"
    # pulls the API from the URL and stores it in a dictionary
    state_park_dict = requests.get(state_park_list_url).json()
    # initializes a list to house all the state park ID numbers
    state_park_list_full = []
    # access the park id value of the state park dictionary and appends them to our list.
    for park in state_park_dict:
        state_park_list_full.append(park['park_id'])
    # initializes a second list for us to store a select number of state park ID numbers in.
    state_park_list_limited = []

    # runs through our full list five times picking different park IDs at random
    for park_id in range(5):
        # creates a variable that is a random number between 0 and the length of our list minus 1
        random_park = random.randint(0, len(state_park_list_full) - 1)
        # checks to see if the park ID number stored in our full list of IDs is already in our limited list.
        while state_park_list_full[random_park] in state_park_list_limited:
            # rolls a new number until it comes back with a park that is not already in our limited list
            # this ensures we won't write information for a park multiple times.
            random_park = random.randint(0, len(state_park_list_full) - 1)
        # adds the state park ID of the random park to our limited list.
        state_park_list_limited.append(state_park_list_full[random_park])

    # Creates the title of our word document.
    state_park_brochure.add_paragraph('Minnesota State Park Guide', 'Title')
    # Loops through each park ID stored in our limited list
    for park_id in state_park_list_limited:
        # calls our mn_state_park_lookup function and runs each park ID through
        # our look up function to pull the unique dictionary assigned of the specific park.
        mn_state_park_lookup(park_id)
        # adds a page break after each park to create a cleaner word document.
        state_park_brochure.add_page_break()
    # saves the word document with the give file name and extension.
    state_park_brochure.save('MN_State_Park_Guide.docx')

# calls our main function
main()
